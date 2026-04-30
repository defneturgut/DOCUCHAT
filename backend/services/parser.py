"""
Doküman ayrıştırma (parsing) servisi.
PDF → PyMuPDF, DOCX → python-docx, TXT → direkt okuma
"""
import io
import re
from pathlib import Path

import fitz          # PyMuPDF
from docx import Document


# Desteklenen uzantılar
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dosya içeriğini düz metne çevirir.
    
    Args:
        file_bytes: Ham dosya içeriği
        filename:   Orijinal dosya adı (uzantı tespiti için)
    
    Returns:
        Temizlenmiş düz metin
    
    Raises:
        ValueError: Desteklenmeyen format veya boş dosya
    """
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Desteklenmeyen dosya formatı: '{suffix}'. "
            f"Kabul edilenler: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if suffix == ".pdf":
        text = _extract_pdf(file_bytes)
    elif suffix in (".docx", ".doc"):
        text = _extract_docx(file_bytes)
    elif suffix == ".txt":
        text = _extract_txt(file_bytes)
    else:
        raise ValueError(f"İşlenemeyen format: {suffix}")

    cleaned = _clean_text(text)

    if len(cleaned.strip()) < 10:
        raise ValueError("Dosya içeriği okunamadı veya çok kısa (< 10 karakter).")

    return cleaned


# ─── Özel Parser'lar ──────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    """PyMuPDF ile PDF'den metin çıkarır."""
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"[Sayfa {page_num}]\n{page_text}")
    return "\n\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    """python-docx ile DOCX/DOC'dan metin çıkarır."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Tablolardan da metin al
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return "\n\n".join(paragraphs)


def _extract_txt(file_bytes: bytes) -> str:
    """TXT dosyasını UTF-8 veya Latin-1 olarak okur."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


# ─── Metin Temizleme ──────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Ham metni temizler:
    - Fazla boşlukları kaldırır
    - Bozuk karakterleri düzeltir
    - Anlamsız satırları atar
    """
    # Satır sonlarını normalize et
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 3'ten fazla art arda boş satırı 2'ye indir
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Satır içi fazla boşlukları temizle
    lines = []
    for line in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", line).strip()
        # Sadece tekrarlayan özel karakterlerden oluşan satırları at
        if not re.match(r"^[=\-_*#~]{3,}$", line):
            lines.append(line)

    return "\n".join(lines).strip()
